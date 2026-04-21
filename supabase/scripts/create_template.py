"""
Genera cotizacion.docx — template mínimo para Carbone.io
Variables basadas en builder.ts del Edge Function generate-pdf
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

BLUE  = RGBColor(0x00, 0x2B, 0x5B)   # Vitali Blue
GOLD  = RGBColor(0xD4, 0xBE, 0x77)   # Vitali Gold
GRAY  = RGBColor(0x66, 0x66, 0x66)
BLACK = RGBColor(0x22, 0x22, 0x22)

def heading(doc, text, size=16, color=BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def row(doc, label, value, label_bold=True):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = label_bold
    r1.font.size = Pt(10)
    r1.font.color.rgb = GRAY
    r2 = p.add_run(value)
    r2.font.size = Pt(10)
    r2.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(2)
    return p

def separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("─" * 72)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ── ENCABEZADO ────────────────────────────────────────────────
p = heading(doc, "VITALI SUITES", size=22, align=WD_ALIGN_PARAGRAPH.CENTER)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Cotización de Departamento")
run.font.size = Pt(13)
run.font.color.rgb = GOLD
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Fecha: {d.cotizacion.fecha}")
run.font.size = Pt(9)
run.font.color.rgb = GRAY

separator(doc)

# ── DATOS DEL CLIENTE ─────────────────────────────────────────
heading(doc, "DATOS DEL CLIENTE", size=12, color=BLUE)
row(doc, "Nombre",   "{d.cliente.nombre}")
row(doc, "RUT",      "{d.cliente.rut}")
row(doc, "Email",    "{d.cliente.email}")
row(doc, "Teléfono", "{d.cliente.phone}")
row(doc, "Dirección","{d.cliente.address}, {d.cliente.commune}")

separator(doc)

# ── DATOS DEL INMUEBLE ────────────────────────────────────────
heading(doc, "INMUEBLE COTIZADO", size=12, color=BLUE)
row(doc, "Proyecto",       "{d.proyecto.nombre}")
row(doc, "Torre",          "{d.torre.nombre}")
row(doc, "Departamento",   "{d.unidad.numero}")
row(doc, "Piso",           "{d.unidad.piso}")
row(doc, "Tipología",      "{d.unidad.tipologia}")
row(doc, "Superficie",     "{d.unidad.superficie} m²")
row(doc, "Fecha entrega",  "{d.torre.fecha_entrega}")
row(doc, "Precio de lista","{d.proyecto.moneda.simbolo} {d.unidad.precio_lista:formatN(0,'.','.')}")

separator(doc)

# ── CONDICIONES DE FINANCIAMIENTO ─────────────────────────────
heading(doc, "CONDICIONES DE FINANCIAMIENTO", size=12, color=BLUE)
row(doc, "Tipo de crédito",    "{d.cotizacion.tipo_credito}")
row(doc, "PIE",                "{d.cotizacion.pie_pct}%  ({d.proyecto.moneda.simbolo} {d.cotizacion.pie_monto:formatN(0,'.','.')})")
row(doc, "Monto a financiar",  "{d.proyecto.moneda.simbolo} {d.cotizacion.monto_credito:formatN(0,'.','.')}")
row(doc, "Plazo",              "{d.cotizacion.plazo_anos} años")
row(doc, "Tasa anual",         "{d.proyecto.tasa_anual}%")

separator(doc)

# ── CRÉDITO FRANCÉS ───────────────────────────────────────────
p = heading(doc, "CRÉDITO FRANCÉS", size=12, color=BLUE)

p2 = doc.add_paragraph()
run = p2.add_run("{d.frances:ifEQ(undefined) hideBegin}")
run.font.size = Pt(1)
run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

row(doc, "Cuota mensual", "{d.frances.cuota_mensual:formatN(2,'.','.')} {d.proyecto.moneda.simbolo}")
row(doc, "Total pagado",  "{d.frances.total_pagado:formatN(2,'.','.')} {d.proyecto.moneda.simbolo}")

p3 = doc.add_paragraph()
run3 = p3.add_run("{d.frances:ifEQ(undefined) hideEnd}")
run3.font.size = Pt(1)
run3.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

separator(doc)

# ── CRÉDITO INTELIGENTE ───────────────────────────────────────
heading(doc, "CRÉDITO INTELIGENTE", size=12, color=BLUE)

p2 = doc.add_paragraph()
run = p2.add_run("{d.inteligente:ifEQ(undefined) hideBegin}")
run.font.size = Pt(1)
run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

row(doc, "Cuota mensual ({d.inteligente.pct_cuotas}% del precio)", "{d.inteligente.cuota_mensual:formatN(2,'.','.')} {d.proyecto.moneda.simbolo}")
row(doc, "Pago final / globo ({d.inteligente.pct_globo}%)",        "{d.inteligente.globo:formatN(2,'.','.')} {d.proyecto.moneda.simbolo}")
row(doc, "Total pagado",                                            "{d.inteligente.total_pagado:formatN(2,'.','.')} {d.proyecto.moneda.simbolo}")

p3 = doc.add_paragraph()
run3 = p3.add_run("{d.inteligente:ifEQ(undefined) hideEnd}")
run3.font.size = Pt(1)
run3.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

separator(doc)

# ── PIE DE PÁGINA ─────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Esta cotización es referencial y no constituye una oferta vinculante.")
run.font.size = Pt(8)
run.font.color.rgb = GRAY
run.italic = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Vitali Suites · contacto@vitalisuites.com · vitalisuites.com")
run2.font.size = Pt(8)
run2.font.color.rgb = GRAY

# ── Guardar ───────────────────────────────────────────────────
out = os.path.join(os.path.dirname(__file__), '..', 'templates', 'cotizacion.docx')
os.makedirs(os.path.dirname(out), exist_ok=True)
doc.save(out)
print(f"Template guardado en: {os.path.abspath(out)}")
